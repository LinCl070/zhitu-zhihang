"""Build the three submission-ready Word documents for 职途智航."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "deliverables"

BLUE = "1F4E79"
INK = "17324D"
MUTED = "5B6573"
LIGHT_BLUE = "EAF1F8"
LIGHT_GRAY = "F3F5F7"
WHITE = "FFFFFF"
TABLE_WIDTH = 9360
TABLE_INDENT = 120


def set_font(run, *, size=11, bold=False, color=INK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    cell_properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell_properties.append(shading)


def set_cell_width(cell, width):
    cell_properties = cell._tc.get_or_add_tcPr()
    width_element = cell_properties.find(qn("w:tcW"))
    if width_element is None:
        width_element = OxmlElement("w:tcW")
        cell_properties.append(width_element)
    width_element.set(qn("w:w"), str(width))
    width_element.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_properties = table._tbl.tblPr
    table_width = table_properties.first_child_found_in("w:tblW")
    table_width.set(qn("w:w"), str(TABLE_WIDTH))
    table_width.set(qn("w:type"), "dxa")
    indent = OxmlElement("w:tblInd")
    indent.set(qn("w:w"), str(TABLE_INDENT))
    indent.set(qn("w:type"), "dxa")
    table_properties.append(indent)
    grid = table._tbl.tblGrid
    for grid_column, width in zip(grid.gridCol_lst, widths, strict=True):
        grid_column.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths, strict=True):
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    cell_properties = cell._tc.get_or_add_tcPr()
    margins = cell_properties.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        cell_properties.append(margins)
    for side, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        element = margins.find(qn(f"w:{side}"))
        if element is None:
            element = OxmlElement(f"w:{side}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instruction)
    run._r.append(end)
    set_font(run, size=9, color=MUTED)


def configure_document(doc, *, running_label):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.2

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 11, INK, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header.add_run(running_label)
    set_font(header_run, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_run = footer.add_run("成都信息工程大学 | 职途智航 | 第 ")
    set_font(footer_run, size=9, color=MUTED)
    add_page_field(footer)
    footer_tail = footer.add_run(" 页")
    set_font(footer_tail, size=9, color=MUTED)


def add_title_block(doc, title, subtitle, metadata):
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(24)
    title_p.paragraph_format.space_after = Pt(8)
    run = title_p.add_run(title)
    set_font(run, size=24, bold=True, color=BLUE)

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_p.paragraph_format.space_after = Pt(18)
    run = subtitle_p.add_run(subtitle)
    set_font(run, size=12, color=MUTED)

    table = doc.add_table(rows=len(metadata), cols=2)
    set_table_geometry(table, [1750, 7610])
    for row, (label, value) in zip(table.rows, metadata, strict=True):
        set_cell_margins(row.cells[0])
        set_cell_margins(row.cells[1])
        set_cell_shading(row.cells[0], LIGHT_GRAY)
        for cell, text, bold in ((row.cells[0], label, True), (row.cells[1], value, False)):
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(text), size=10, bold=bold)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_paragraph(doc, text, *, emphasis=None):
    paragraph = doc.add_paragraph()
    if emphasis:
        label, content = emphasis
        set_font(paragraph.add_run(label), bold=True, color=BLUE)
        set_font(paragraph.add_run(content))
    else:
        set_font(paragraph.add_run(text))
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    header = table.rows[0]
    for cell, text in zip(header.cells, headers, strict=True):
        set_cell_margins(cell)
        set_cell_shading(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(text), size=9.5, bold=True, color=BLUE)
    for row_values in rows:
        cells = table.add_row().cells
        for cell, text in zip(cells, row_values, strict=True):
            set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_font(p.add_run(text), size=9.2)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_case_plan():
    doc = Document()
    configure_document(doc, running_label="成果应用案例方案")
    add_title_block(
        doc,
        "职途智航成果应用案例方案",
        "成都信息工程大学就业择业规划智能体",
        [
            ("参赛方向", "校园管理方向"),
            ("适用场景", "成都信息工程大学就业指导与学生发展场景"),
            ("版本", "v1.0"),
            ("日期", "2026-07-23"),
        ],
    )
    doc.add_heading("一、应用背景与痛点", level=1)
    add_paragraph(doc, "高校毕业生在求职准备阶段常面临岗位信息分散、职业目标模糊、能力差距难量化、就业政策难检索等问题。就业指导教师需要为不同专业、不同求职阶段的学生提供建议，但人工咨询时间有限，且建议过程难以持续追踪。")
    add_paragraph(doc, "现有通用问答工具难以保证岗位来源、推荐依据和隐私边界。职途智航聚焦在不收集非必要个人信息的前提下，为学生提供可解释岗位匹配、可执行行动计划和可信资料咨询。")

    doc.add_heading("二、解决方案", level=1)
    add_paragraph(doc, "系统以学生主动填写的脱敏职业画像为输入，结合已审核的岗位与就业资料，形成“画像 - 匹配 - 差距 - 计划 - 咨询 - 教师协助”的闭环。")
    add_table(
        doc,
        ["功能", "用户价值", "实现方式"],
        [
            ("职业画像", "梳理专业、技能、项目、目标岗位和城市", "受控表单与字段校验，只保存规划所需信息"),
            ("岗位推荐", "理解岗位被推荐的原因", "固定规则计算分数并展示分项、缺口和来源"),
            ("行动计划", "将能力差距转为阶段性任务", "基于匹配缺口生成待办，不承诺录用结果"),
            ("智能咨询", "查询职业准备与就业政策资料", "SF-FastGPT 工作流检索，无可靠来源时人工确认"),
            ("教师协助", "在授权范围内获得持续指导", "展示职业摘要、计划状态和教师本人意见"),
            ("管理员工作台", "维护资料与授权边界", "展示审核追溯与脱敏审计，不展示密钥或完整画像"),
        ],
        [1450, 3150, 4760],
    )

    doc.add_heading("三、创新点与可推广性", level=1)
    for label, content in [
        ("可解释而非黑箱推荐。", "岗位匹配固定采用技能 50 分、专业 20 分、项目 15 分、城市 5 分、目标岗位 10 分的规则权重，学生可看到分项得分、能力缺口和信息来源。"),
        ("知识检索与业务数据分离。", "PostgreSQL 保存最小化业务数据，SF-FastGPT 仅负责受审核资料检索与工作流，平台不直接访问数据库。"),
        ("安全优先的人工兜底。", "平台异常、资料缺失或回答没有可引用来源时，系统统一给出人工确认建议，不伪造政策、岗位或录用结论。"),
        ("三角色协同。", "学生、就业指导教师和管理员具有不同数据可见范围，适合就业指导部门的实际协作过程。"),
    ]:
        add_paragraph(doc, "", emphasis=(label, content))

    doc.add_heading("四、数据与知识来源", level=1)
    add_table(
        doc,
        ["数据类别", "当前来源", "使用边界"],
        [
            ("岗位数据", "项目组构造的模拟岗位台账", "仅用于本地验证与演示，不表述为真实招聘信息"),
            ("就业政策与职业资料", "来源登记后审核", "须具备来源、日期、适用范围和审核状态后才可上传知识库"),
            ("学生职业画像", "学生主动输入的脱敏信息", "不采集联系方式、学号、性别、籍贯、健康、政治面貌等信息"),
            ("教师意见与审计", "系统受控写入", "仅限授权范围；审计不保存请求体、画像、令牌或密钥"),
        ],
        [1450, 2600, 5310],
    )

    doc.add_heading("五、预期应用效果", level=1)
    add_table(
        doc,
        ["目标", "可观察指标", "验证方式"],
        [
            ("提升规划清晰度", "学生完成画像、查看缺口并生成计划", "学生端完整闭环演示"),
            ("提升建议可解释性", "每个推荐展示总分、分项、缺口和来源", "匹配结果页面与接口验证"),
            ("降低重复答疑负担", "教师只查看授权学生摘要并提交定向意见", "教师工作台演示"),
            ("控制隐私风险", "敏感字段与密钥不出现在前端、日志或演示画面", "自动化隐私检查与脱敏审计"),
        ],
        [1750, 4500, 3110],
    )
    add_paragraph(doc, "首版仅使用模拟数据，未对真实学生就业结果作出统计结论。正式试点前，应由学校就业指导部门审核知识资料、岗位来源、数据授权和使用流程。")

    doc.add_heading("六、实施范围与后续计划", level=1)
    add_paragraph(doc, "首版已实现本地 PostgreSQL 数据持久化、三角色权限控制、确定性匹配、行动计划、模拟咨询模式、教师意见、管理员资料与审计工作台。赛事平台提供基础地址、服务端 API Key 和两个工作流 ID 后，可完成 SF-FastGPT 真实知识库联调。")
    add_paragraph(doc, "不在首版范围内的事项包括接入学校生产数据库、接入真实学生账号、自动同步招聘平台，以及对就业结果作出预测或承诺。")
    doc.save(OUTPUT / "01-成果应用案例方案.docx")


def build_technical_document():
    doc = Document()
    configure_document(doc, running_label="成果实施技术文档")
    add_title_block(
        doc,
        "职途智航成果实施技术文档",
        "成都信息工程大学就业择业规划智能体",
        [("版本", "v1.0"), ("部署范围", "本地 PostgreSQL 脱敏演示环境"), ("日期", "2026-07-23")],
    )
    doc.add_heading("一、技术目标", level=1)
    add_paragraph(doc, "项目实现一个可本地部署、可演示、可审计的高校就业择业规划智能体。系统支持学生职业规划闭环、教师授权协助和管理员资料治理，同时避免将数据库凭据、平台密钥或非必要学生信息暴露给浏览器与模型平台。")

    doc.add_heading("二、总体架构", level=1)
    add_table(
        doc,
        ["层级", "采用技术", "职责"],
        [
            ("前端", "React、TypeScript、Vite", "学生、教师、管理员工作台与响应式界面"),
            ("后端", "FastAPI、Pydantic", "接口、输入校验、权限控制和安全响应"),
            ("数据访问", "SQLAlchemy 2、Alembic、psycopg", "PostgreSQL 模型、迁移、事务和索引"),
            ("智能体适配", "SF-FastGPT Chat Completions", "工作流路由、资料来源归一化和安全降级"),
            ("数据治理", "CSV 台账与 Markdown 规范", "模拟岗位、资料来源、审核状态和使用边界"),
        ],
        [1550, 3050, 4760],
    )

    doc.add_heading("三、关键实现", level=1)
    doc.add_heading("3.1 三角色权限", level=2)
    add_paragraph(doc, "学生只能读写自己的职业画像、匹配结果和行动计划。教师需要管理员显式授权才能查看某位学生的职业摘要、计划状态和本人意见。管理员可维护授权并查看岗位、资料审核字段以及脱敏审计，但不能读取完整画像、意见正文、数据库连接串或 FastGPT 密钥。")
    add_paragraph(doc, "演示登录使用服务端内存会话，不保存密码；令牌重启后失效，最长有效期为 8 小时。所有越权访问都记录为仅含角色、操作、资源类型、资源标识和时间的脱敏审计事件。")
    doc.add_heading("3.2 可解释岗位匹配", level=2)
    add_paragraph(doc, "岗位匹配为不依赖大模型的确定性服务，只从已发布、未过期且有来源标题的岗位中选择。总分为 100 分：技能 50 分、专业 20 分、项目 15 分、目标城市 5 分、目标岗位 10 分。服务返回分项得分、未满足技能、岗位来源与稳定排序。")
    doc.add_heading("3.3 SF-FastGPT 工作流", level=2)
    add_paragraph(doc, "后端为职业咨询与政策咨询定义统一响应契约，包含回答、来源、免责声明、人工兜底标记和运行模式。职业与政策工作流分别映射到服务端环境变量中的工作流 ID，前端不保存平台地址、密钥或工作流 ID。")
    add_paragraph(doc, "当运行在 mock 模式时，系统返回统一的本地人工确认建议。切换为真实模式后，服务端向可配置的 Chat Completions 路径发送最小化问题与上下文。平台请求失败、格式异常或没有可信来源时，系统转为人工确认。")
    doc.add_heading("3.4 数据与隐私边界", level=2)
    add_table(
        doc,
        ["边界", "控制措施"],
        [
            ("最小化采集", "不保存或推断电话、邮箱、学号、性别、籍贯、健康、政治面貌、家庭情况和完整成绩单。"),
            ("资料可信", "岗位与知识资料必须具备来源、日期、状态和适用范围；无来源资料不能进入推荐。"),
            ("密钥隔离", "数据库 URL、API Key 与工作流 ID 仅存在于 .env 或部署环境。"),
            ("演示边界", "当前岗位为模拟数据，不得作为真实招聘承诺或上传为真实招聘知识。"),
        ],
        [2100, 7260],
    )

    doc.add_heading("四、部署与运行", level=1)
    add_paragraph(doc, "前置条件包括 Node.js、PostgreSQL 14+、位于 backend/.venv 的 Python 虚拟环境，以及根目录 .env 中已配置的 PostgreSQL DATABASE_URL。真实 SF-FastGPT 联调另需平台地址、服务端 API Key 和两个工作流 ID。")
    add_table(
        doc,
        ["步骤", "命令或操作"],
        [
            ("数据库迁移", "设置 PYTHONPATH=backend 后执行 alembic upgrade head。"),
            ("启动后端", "使用 uvicorn 启动 app.main:app，端口 8000。"),
            ("启动前端", "在 frontend 目录执行 npm run dev -- --port 5174。"),
            ("真实联调", "配置平台参数后运行 python -m scripts.verify_fastgpt。"),
        ],
        [2200, 7160],
    )
    add_paragraph(doc, "联调脚本只输出职业与政策两个工作流的运行模式、来源数量、来源完整性和通过状态。两条工作流均应返回带标题、链接、版本或日期的资料来源，且不得触发人工兜底。")

    doc.add_heading("五、测试与验证", level=1)
    add_table(
        doc,
        ["验证项", "当前结果"],
        [
            ("后端单元测试", "26 项通过，覆盖配置、模型、RBAC、匹配、FastGPT 适配、隐私、教师意见与联调诊断。"),
            ("后端静态检查", "Ruff 检查和格式校验通过。"),
            ("前端质量检查", "ESLint 通过，TypeScript 与 Vite 生产构建成功。"),
            ("数据库迁移", "当前版本为 20260723_0003 (head)。"),
            ("页面验证", "三角色关键路径已在本地演示；学生与教师移动端 375px 无水平溢出。"),
            ("真实 SF-FastGPT 联调", "待赛事平台提供并配置 API Key 与两个工作流 ID 后执行。"),
        ],
        [2600, 6760],
    )
    doc.add_heading("六、运维与扩展建议", level=1)
    add_paragraph(doc, "正式试点前应由就业指导部门审核资料来源、更新周期和适用范围，并建立资料失效机制。后续可增加学院维度资料隔离、更多行动计划模板、人工咨询预约及审批流。任何生产数据库、真实学生身份认证或外部招聘数据同步接入，都应重新进行数据授权与安全评审。")
    doc.save(OUTPUT / "02-成果实施技术文档.docx")


def build_video_script():
    doc = Document()
    configure_document(doc, running_label="成果演示视频脚本")
    add_title_block(
        doc,
        "职途智航成果演示视频脚本",
        "建议时长：5 分 30 秒",
        [("演示环境", "本地 PostgreSQL + 脱敏演示账号 + 模拟岗位数据"), ("日期", "2026-07-23")],
    )
    doc.add_heading("一、录制前检查", level=1)
    for label, content in [
        ("画面边界：", "仅展示本地演示地址，不展示 .env、数据库连接串、API Key、工作流 ID、终端密码或真实学生信息。"),
        ("数据边界：", "使用合成演示账号和模拟岗位，保留“演示数据”语义。"),
        ("运行检查：", "先启动 PostgreSQL、后端和前端；确认学生、教师、管理员三角色均可登录。"),
        ("平台状态：", "真实 FastGPT 未配置时，展示模拟模式的人工兜底，不宣称真实平台联调已完成。"),
    ]:
        add_paragraph(doc, "", emphasis=(label, content))
    doc.add_heading("二、分镜与旁白", level=1)
    add_table(
        doc,
        ["时间", "画面操作", "旁白要点"],
        [
            ("00:00-00:20", "打开首页，展示项目与学校品牌", "说明项目面向高校学生就业择业规划，形成可解释、可追溯的闭环。"),
            ("00:20-00:45", "展示三种演示入口", "说明学生、教师、管理员三角色协同。"),
            ("00:45-01:25", "学生登录并填写职业画像", "强调仅使用规划所需脱敏信息，不采集非必要个人信息。"),
            ("01:25-02:05", "展示岗位推荐、分数与能力缺口", "说明固定规则及五项分数来源，展示可解释推荐。"),
            ("02:05-02:35", "生成行动计划", "说明系统将能力缺口转换为阶段性待办，不作录用承诺。"),
            ("02:35-03:10", "提问职业准备或政策问题", "说明咨询通过服务端适配 FastGPT，前端不保存密钥。"),
            ("03:10-03:35", "模拟模式展示人工确认建议", "资料不足、无来源或平台不可用时不编造答案，建议人工确认。"),
            ("03:35-04:10", "管理员查看台账、授权与审计", "说明管理员只处理审核追溯字段和授权边界。"),
            ("04:10-04:40", "教师查看授权学生并提交意见", "说明教师只看到职业摘要、计划状态和本人意见。"),
            ("04:40-05:05", "展示移动端学生与教师视图", "说明 375px 宽度下无水平溢出。"),
            ("05:05-05:30", "回到首页或架构图", "总结智能体结合匹配、检索、人工兜底与权限治理。"),
        ],
        [1350, 3250, 4760],
    )
    doc.add_heading("三、真实 SF-FastGPT 联调补录段", level=1)
    add_paragraph(doc, "仅在赛事平台参数完成配置并运行验收脚本通过后，替换 03:10-03:35 段。画面展示咨询回答中的资料标题、链接与版本或日期，不展示开发者工具、请求头、.env 或终端环境变量。")
    doc.add_heading("四、提交检查清单", level=1)
    add_table(
        doc,
        ["检查项", "要求"],
        [
            ("品牌信息", "视频中出现项目名称和“成都信息工程大学”。"),
            ("功能覆盖", "覆盖学生画像、可解释匹配、行动计划、咨询、教师协助、资料治理和审计。"),
            ("信息安全", "不出现真实学生资料、联系方式、密码、数据库 URL、API Key 或工作流 ID。"),
            ("演示真实性", "明确模拟岗位边界，不作真实招聘或录用承诺。"),
        ],
        [2000, 7360],
    )
    doc.save(OUTPUT / "03-成果演示视频脚本.docx")


if __name__ == "__main__":
    OUTPUT.mkdir(parents=True, exist_ok=True)
    build_case_plan()
    build_technical_document()
    build_video_script()
